"""
Handwriting and Math Exam OCR Engine using Hugging Face Qwen-VL Model

This script provides a complete OCR system for handwritten English text and mathematical equations.
It uses the Qwen-VL vision-language model to transcribe images and parse exam answers.

Requirements:
- Python 3.x
- Virtual environment (venv) recommended
- GPU support for faster processing (optional)

Author: GitHub Copilot
Date: October 20, 2025
"""

import os
import json
import subprocess
import sys
from pathlib import Path
from typing import Dict, List, Optional, Union

import torch
from transformers import Qwen2VLForConditionalGeneration, AutoTokenizer, Qwen2VLProcessor
import accelerate
from PIL import Image
import matplotlib.pyplot as plt
from docx import Document


# Package installation function
def install_packages():
    """Automatically install required packages if not present."""
    required_packages = [
        'torch', 'torchvision', 'transformers', 'accelerate', 'huggingface_hub', 'pillow', 'matplotlib', 'python-docx', 'opencv-python'
    ]
    missing_packages = []

    for package in required_packages:
        try:
            if package == 'opencv-python':
                __import__('cv2')
            else:
                __import__(package.replace('-', '_'))
        except ImportError:
            missing_packages.append(package)

    if missing_packages:
        print(f"Installing missing packages: {', '.join(missing_packages)}")
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install'] + missing_packages)
            print("Packages installed successfully.")
        except subprocess.CalledProcessError as e:
            print(f"Failed to install packages: {e}")
            sys.exit(1)


# Install packages at startup
install_packages()


class OCREngine:
    """OCR Engine using Qwen-VL model for handwritten text and math equations."""

    def __init__(self, model_name: str = "Qwen/Qwen2-VL-2B-Instruct"):
        """
        Initialize the OCR engine with the specified model.

        Args:
            model_name: Hugging Face model name for Qwen-VL
        """
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        print(f"Using device: {self.device}")

        try:
            self.model = Qwen2VLForConditionalGeneration.from_pretrained(
                model_name,
                torch_dtype=torch.float16 if self.device.type == "cuda" else torch.float32,
                device_map="auto" if self.device.type == "cuda" else None
            )
            self.tokenizer = AutoTokenizer.from_pretrained(model_name)
            self.processor = Qwen2VLProcessor.from_pretrained(model_name, use_fast=False)
            print("Model loaded successfully.")
        except Exception as e:
            print(f"Error loading model: {e}")
            print("Make sure you have access to the model and sufficient resources.")
            sys.exit(1)

    def ocr_handwriting(self, image_path: str) -> str:
        """
        Perform OCR on a single image containing handwritten text.

        Args:
            image_path: Path to the image file

        Returns:
            Raw OCR text transcription
        """
        try:
            # Load and preprocess image
            image = Image.open(image_path).convert("RGB")

            # Create prompt for OCR
            prompt = (
                """
<|system|>
You are a high-accuracy OCR and layout transcription assistant designed for hybrid exam sheets. 
Your task is to reproduce the *exact visible text, layout, and structure* of the input image. 
You must behave like a scanner that outputs text faithfully — not like a language model that interprets meaning.

Key principle: **Every visible mark, word, or symbol matters.**
You must never rephrase, summarize, correct spelling, or omit visual elements.

You are transcribing an exam answer sheet that may contain a mix of printed questions and handwritten answers. 
Your job is to extract everything visible — printed or handwritten — exactly as it appears. 
Follow these detailed rules strictly:

-----------------------------------------------------
🔹 1. TEXT FIDELITY & SYMBOL PRESERVATION
-----------------------------------------------------
- Capture all visible text exactly as it is written.
- DO NOT rewrite, correct grammar, or infer missing text.
- Preserve all visible question numbers such as:
  Q1. / Q2) / Q3 : / (a) / (i) / (1) etc.
- Include all symbols:  
  ✯ ★ * • → ← ↑ ↓ = + − × ÷ ± ~ ° α β π θ ∞ ≤ ≥ √ ∫ ∑  
  and any other mathematical or decorative marks.
- Retain punctuation and special characters (colons, brackets, commas, etc.).
- If the writer used unique markers (like a custom star or dash), transcribe it literally.
- If any symbol or word is unreadable, write `[illegible]`.

-----------------------------------------------------
🔹 2. LINE STRUCTURE & SPACING
-----------------------------------------------------
- Every new handwritten or printed line in the image = new line in your output.
- Do not merge separate lines into a single sentence.
- Preserve blank lines between question sections (add one empty line between Q1 and Q2, etc.).
- Maintain indentation if the text is indented in the image.
- Do not automatically add or remove spacing — mirror what’s visible.

-----------------------------------------------------
🔹 3. QUESTION SEPARATION LOGIC
-----------------------------------------------------
- Identify each question block clearly.
- A question block usually starts with Q1, Q2, Q3, etc.
- Each question block must begin on a new line, formatted like this:
Q1.
[content]

csharp
Copy code
- If the question and answer appear together, label as:
Q1.
Question (Printed): ...
Answer (Handwritten): ...

vbnet
Copy code
- If there’s no printed question (just answers), still separate them based on visible labels (Q1, Q2, etc.).

-----------------------------------------------------
🔹 4. BULLETS, STARS, CHECKMARKS, & LIST ITEMS
-----------------------------------------------------
- Retain all bullet types exactly as written:
- • small filled circle → output as "•"
- ✯ or ★ star → output as "✯" or "★"
- * asterisk → output as "*"
- → arrow → output as "→"
- Keep bullet alignment (each bullet on a new line).
- Do not remove or standardize bullets. Each must remain exactly as seen.

-----------------------------------------------------
🔹 5. MATHEMATICAL CONTENT & EQUATIONS
-----------------------------------------------------
- Represent all math and scientific notation in plain Unicode.
Examples:
- Use “²” and “³” for powers (not ^2 or ^3)
- Write fractions as “½”, “¾”, “⅓”, etc. if visible.
- Preserve equal signs (=), plus/minus (±), multiply (×), divide (÷).
- If the equation is handwritten, keep spacing identical:
Example: “a + b² + 2d = 7g”
- Never reformat math into LaTeX — keep it in simple text.

-----------------------------------------------------
🔹 6. HEADERS, FOOTERS, & NON-TEXT ELEMENTS
-----------------------------------------------------
- If visible, include metadata like:
- Student name, Roll No., Subject, Date.
- Page numbers or exam titles.
- If there’s a diagram, chart, or arrow drawing:
Write: `[Diagram: <brief description>]`
Example: `[Diagram: a labeled triangle with sides a, b, c]`
- If lines, boxes, or tables appear, describe them briefly in brackets:
`[Table: 2 rows, 3 columns]`

-----------------------------------------------------
🔹 7. UNCLEAR REGIONS
-----------------------------------------------------
- For smudged or unclear handwriting, use `[illegible]`.
- Do not guess what the text might be.

-----------------------------------------------------
🔹 8. OUTPUT STRUCTURE FORMAT
-----------------------------------------------------
Produce output in this exact clean format:

Q1.
✯ First Point
✯ Second Point

Q2)
• First Point
• Second Point

Q3
Hi

Q4:
a + b² + 2d = 7g

pgsql
Copy code

-----------------------------------------------------
🔹 9. DO NOTS
-----------------------------------------------------
- ❌ Do NOT summarize.
- ❌ Do NOT reorder questions.
- ❌ Do NOT fill missing text.
- ❌ Do NOT change bullet types or symbols.
- ❌ Do NOT correct spelling or grammar.
- ❌ Do NOT output explanations or commentary.

-----------------------------------------------------
🔹 10. OUTPUT MODE
-----------------------------------------------------
- Output only clean text transcription — no markdown or metadata.
- Wrap code blocks (```) around the full output ONLY if the system requires monospaced formatting.

"""
            )

            # Prepare inputs
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image", "image": image}
                    ]
                }
            ]

            # Process inputs
            text = self.processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
            inputs = self.processor(text=text, images=image, return_tensors="pt").to(self.device)

            # Generate response
            with torch.no_grad():
                generated_ids = self.model.generate(
                    **inputs,
                    max_new_tokens=2048,
                    do_sample=False,
                    temperature=0.0
                )

            # Decode output
            generated_ids_trimmed = [
                out_ids[len(in_ids):] for in_ids, out_ids in zip(inputs.input_ids, generated_ids)
            ]
            output_text = self.processor.batch_decode(
                generated_ids_trimmed, skip_special_tokens=True, clean_up_tokenization_spaces=False
            )[0]

            return output_text.strip()

        except Exception as e:
            print(f"Error processing image {image_path}: {e}")
            return ""

    def parse_exam_output(self, ocr_text: str) -> Dict[str, str]:
        """
        Parse raw OCR text into structured exam answers.

        Args:
            ocr_text: Raw OCR output text

        Returns:
            Dictionary with question numbers as keys and answers as values
        """
        structured_answers = {}

        # Split by lines and look for question patterns
        lines = ocr_text.split('\n')
        current_question = None
        current_answer = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if line starts with question number (e.g., "1.", "Question 1:", etc.)
            import re
            question_match = re.match(r'^(?:\(?\w+\)?\s*)?(\d+)\)?\.?\s*(.*)$', line, re.IGNORECASE)
            if question_match:
                # Save previous question if exists
                if current_question is not None:
                    structured_answers[str(current_question)] = ' '.join(current_answer).strip()

                current_question = int(question_match.group(1))
                current_answer = [question_match.group(2)]
            else:
                # Continue accumulating answer
                if current_question is not None:
                    current_answer.append(line)

        # Save last question
        if current_question is not None:
            structured_answers[str(current_question)] = ' '.join(current_answer).strip()

        return structured_answers

    def process_folder(self, folder_path: str) -> Dict[str, Dict[str, str]]:
        """
        Process all images in a folder and return structured results.

        Args:
            folder_path: Path to folder containing images

        Returns:
            Dictionary with image filenames as keys and structured answers as values
        """
        results = {}
        folder = Path(folder_path)

        if not folder.exists() or not folder.is_dir():
            print(f"Folder {folder_path} does not exist or is not a directory.")
            return results

        # Supported image extensions
        image_extensions = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}

        for file_path in folder.iterdir():
            if file_path.suffix.lower() in image_extensions:
                print(f"Processing {file_path.name}...")
                ocr_text = self.ocr_handwriting(str(file_path))
                print(f"Raw OCR for {file_path.name}: {ocr_text}")  # Debug: print raw OCR text
                structured = self.parse_exam_output(ocr_text) if ocr_text else {}
                results[file_path.name] = {
                    'raw_ocr': ocr_text,
                    'structured': structured
                }

        return results

    def save_results(self, results: Dict[str, Dict[str, Union[str, Dict[str, str]]]], output_path: str, format: str = 'json'):
        """
        Save OCR results to file.

        Args:
            results: OCR results with raw and structured data
            output_path: Path to output file
            format: Output format ('json', 'docx', 'txt')
        """
        if format.lower() == 'json':
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"Results saved to {output_path} as JSON.")

        elif format.lower() == 'docx':
            doc = Document()
            doc.add_heading('Exam OCR Results', 0)

            for image_name, data in results.items():
                doc.add_heading(f'Image: {image_name}', level=1)
                raw_ocr = data.get('raw_ocr', '')
                structured = data.get('structured', {})
                
                doc.add_heading('Raw OCR', level=2)
                doc.add_paragraph(raw_ocr)
                
                doc.add_heading('Structured Answers', level=2)
                if structured:
                    for q_num, answer in structured.items():
                        doc.add_paragraph(f'Question {q_num}: {answer}')
                else:
                    doc.add_paragraph('No answers extracted.')

            doc.save(output_path)
            print(f"Results saved to {output_path} as DOCX.")

        elif format.lower() == 'txt':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("Exam OCR Results\n\n")
                for image_name, data in results.items():
                    f.write(f"Image: {image_name}\n")
                    raw_ocr = data.get('raw_ocr', '')
                    structured = data.get('structured', {})
                    
                    f.write("Raw OCR:\n")
                    f.write(raw_ocr + "\n\n")
                    
                    f.write("Structured Answers:\n")
                    if structured:
                        for q_num, answer in structured.items():
                            f.write(f"Question {q_num}: {answer}\n")
                    else:
                        f.write("No answers extracted.\n")
                    f.write("\n")
            print(f"Results saved to {output_path} as TXT.")

        else:
            print(f"Unsupported format: {format}. Use 'json', 'docx', or 'txt'.")


def main():
    """Main demonstration script."""
    print("Initializing OCR Engine...")

    # Initialize engine
    engine = OCREngine()

    # Example usage
    print("\n=== Demonstration ===")

    # Create sample folder if it doesn't exist
    sample_folder = Path("sample_images")
    sample_folder.mkdir(exist_ok=True)

    # Note: In a real scenario, place your images in sample_images folder
    print(f"Place your handwritten exam images in the '{sample_folder}' folder.")

    # Process folder
    if sample_folder.exists() and any(sample_folder.iterdir()):
        results = engine.process_folder(str(sample_folder))

        # Save results
        engine.save_results(results, "ocr_results.json", format='json')
        engine.save_results(results, "ocr_results.docx", format='docx')
        engine.save_results(results, "ocr_results.txt", format='txt')

        print("\nProcessing complete. Check the output files.")
    else:
        print("No images found in sample_images folder. Add some images to test.")

    # Single image example (uncomment and modify path as needed)
    # image_path = "path/to/your/image.jpg"
    # ocr_text = engine.ocr_handwriting(image_path)
    # structured = engine.parse_exam_output(ocr_text)
    # print("OCR Text:", ocr_text)
    # print("Structured:", structured)


if __name__ == "__main__":
    main()